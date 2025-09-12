# generators/docx_generator.py
# -*- coding: utf-8 -*-
"""
Erstellt Berichtshefte im DOCX-Format, basierend auf einer Textvorlage.
"""
import logging
from typing import Dict, Any
from docx import Document
from docx.shared import Pt

from generators.base_generator import BaseGenerator
from core import config

logger = logging.getLogger(__name__)

class DocxGenerator(BaseGenerator):
    """
    Spezialisierte Klasse zur Generierung von DOCX-Berichtsheften.
    Das Layout orientiert sich an einem einfachen Textformat.
    """
    def __init__(self, context: Dict[str, Any]):
        super().__init__(context)
        self.doc: Document = None

    def _setup_document(self) -> None:
        """Initialisiert ein neues DOCX-Dokument."""
        self.doc = Document()
        # Grundlegende Stileinstellungen für das gesamte Dokument
        style = self.doc.styles['Normal']
        font = style.font
        font.name = config.DOCX_FONT_BODY
        font.size = Pt(11)

    def _create_header(self) -> None:
        """Erstellt die Kopfzeile des DOCX-Dokuments."""
        # Titelzeile
        p_titel = self.doc.add_paragraph()
        run_titel = p_titel.add_run(f'Ausbildungsnachweis Nr. {self.context.get("fortlaufende_nr", "")}')
        run_titel.font.name = config.DOCX_FONT_HEADLINE
        run_titel.font.size = Pt(20)
        
        # Inhaltszeile
        azubi = self.context.get("name_azubi", "")
        zeitraum_von = self.context.get("zeitraum_von", "")
        zeitraum_bis = self.context.get("zeitraum_bis", "")
        aj = self.context.get("ausbildungsjahr", "")
        
        p_info = self.doc.add_paragraph(f'Azubi: {azubi}; Zeitraum: {zeitraum_von} bis {zeitraum_bis}; Jahr {aj}')
        p_info.runs[0].font.size = Pt(11)
        # self.doc.add_paragraph() # Leerer Absatz für Abstand

    def _create_body(self) -> None:
        """Erstellt den Hauptteil mit den täglichen Berichtsdaten als Textblöcke."""
        tage_daten = self.context.get("tage_daten", [])
        for i, tag_name in enumerate(config.WOCHENTAGE):
            tag_daten = tage_daten[i] if i < len(tage_daten) else {}
            
            # Info-Zeile für den Tag
            info_zeile = self.doc.add_paragraph()
            run_info = info_zeile.add_run(f'{tag_name}; Typ: {tag_daten.get("typ", "")}; Gesamtstunden: {tag_daten.get("stunden", "")}')
            run_info.font.name = config.DOCX_FONT_HEADLINE
            run_info.font.size = Pt(12)
            
            # Tätigkeiten als Aufzählungspunkte
            taetigkeiten = tag_daten.get("taetigkeiten", "-").split('\n')
            for item in taetigkeiten:
                if item.strip(): # Nur hinzufügen, wenn Zeile nicht leer ist
                    self.doc.add_paragraph(item.strip(), style='List Bullet')

            if i < len(config.WOCHENTAGE) - 1:
                pass # oder  self.doc.add_paragraph() # Abstand zwischen den Tagen


    def _create_footer(self) -> None:
        """Erstellt die Fußzeile mit Datum und Unterschriftsfeldern."""
        # self.doc.add_paragraph() # Abstand
        
        datum_azubi = self.context.get("erstellungsdatum_bericht", "")

        # Auszubildender
        p_azubi = self.doc.add_paragraph()
        run_azubi = p_azubi.add_run('Auszubildender')
        run_azubi.font.name = config.DOCX_FONT_HEADLINE
        run_azubi.font.size = Pt(12)

        p_azubi_datum = self.doc.add_paragraph(f"Datum: {datum_azubi}; Unterschrift:")
        p_azubi_datum.runs[0].font.size = Pt(11)

        # self.doc.add_paragraph()

        # Ausbilder
        p_ausbilder = self.doc.add_paragraph()
        run_ausbilder = p_ausbilder.add_run('Ausbildender bzw. Ausbilder')
        run_ausbilder.font.name = config.DOCX_FONT_HEADLINE
        run_ausbilder.font.size = Pt(12)

        p_ausbilder_datum = self.doc.add_paragraph("Datum: .................; Unterschrift:")
        p_ausbilder_datum.runs[0].font.size = Pt(11)


    def _save_document(self, dateiname: str) -> None:
        """Speichert das DOCX-Dokument."""
        try:
            self.doc.save(dateiname)
            logger.info(f"DOCX-Dokument erfolgreich gespeichert: {dateiname}")
        except Exception as e:
            logger.error(f"Fehler beim Speichern des DOCX-Dokuments '{dateiname}'.", exc_info=True)
            raise IOError(f"Konnte DOCX nicht speichern: {e}")