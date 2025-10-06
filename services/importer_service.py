# services/importer_service.py
# -*- coding: utf-8 -*-
"""
Dienst zur Kapselung der Logik für den Import von DOCX-Berichtsheften.
"""

import logging
import re
import os
from typing import Dict, Optional, Any, List, Tuple
from docx import Document
from datetime import datetime

from core import config

logger = logging.getLogger(__name__)

class ImporterService:
    """
    Liest Daten aus einem vorhandenen DOCX-Berichtsheft und wandelt sie in ein strukturiertes Format um.
    """
    def _extract_from_regex(self, text: str, pattern: str) -> Optional[str]:
        """Hilfsfunktion, um einen Wert mit Regex aus einem Text zu extrahieren."""
        match = re.search(pattern, text, re.IGNORECASE) # Ignoriert Groß/Kleinschreibung
        return match.group(1).strip() if match else None

    def parse_docx(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:
        """
        Analysiert eine einzelne DOCX-Datei und extrahiert die Berichtsdaten.

        Args:
            file_path: Der Pfad zur DOCX-Datei.

        Returns:
            Ein Tupel (context, error_message). Bei Erfolg ist error_message None.
        """
        if os.path.basename(file_path).startswith('~$'):
            logger.debug(f"Überspringe temporäre Word-Datei: {os.path.basename(file_path)}")
            return None, "Temporäre Datei übersprungen."
            
        try:
            doc = Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])

            context: Dict[str, Any] = {}

            # Robustere Regex-Muster, die mehr Variationen zulassen
            context["fortlaufende_nr"] = self._extract_from_regex(full_text, r"Ausbildungsnachweis\s*Nr\.?\s*(\d+)")
            context["name_azubi"] = self._extract_from_regex(full_text, r"Azubi:\s*(.*?);")
            zeitraum_von_str = self._extract_from_regex(full_text, r"Zeitraum:\s*(\d{2}\.\d{2}\.\d{4})")
            
            if not context["fortlaufende_nr"]:
                return None, "Konnte 'Berichtsnummer' nicht finden."
            if not context["name_azubi"]:
                return None, "Konnte 'Name des Azubis' nicht finden."
            if not zeitraum_von_str:
                return None, "Konnte 'Zeitraum' nicht finden."

            start_datum = datetime.strptime(zeitraum_von_str, "%d.%m.%Y").date()
            jahr, kw, _ = start_datum.isocalendar()
            context["jahr"] = str(jahr)
            context["kalenderwoche"] = str(kw)
            context["fortlaufende_nr"] = int(context["fortlaufende_nr"])

            tage_daten: List[Dict[str, str]] = []
            
            # Verbesserte Logik zur Erkennung von Tagesabschnitten
            # Wir suchen nach einem Wochentag am Anfang eines Paragraphen
            day_pattern = re.compile(r"^\s*(" + "|".join(config.DAYS_IN_WEEK) + r")\s*;", re.IGNORECASE)
            
            current_day_data = {}
            for para in doc.paragraphs:
                match = day_pattern.match(para.text)
                if match:
                    # Wenn ein neuer Tag gefunden wird, den vorherigen speichern
                    if current_day_data:
                        tage_daten.append(current_day_data)
                    
                    # Neuen Tag initialisieren
                    current_day_data = {
                        "tag_name": match.group(1).capitalize(),
                        "typ": self._extract_from_regex(para.text, r"Typ:\s*(\w+)") or "Betrieb",
                        "stunden": self._extract_from_regex(para.text, r"Gesamtstunden:\s*([\d:]+)") or "0:00",
                        "taetigkeiten": ""
                    }
                elif current_day_data:
                    # Text zum aktuellen Tag hinzufügen
                    if para.text.strip():
                        # Behält Zeilenumbrüche bei
                        if current_day_data["taetigkeiten"]:
                            current_day_data["taetigkeiten"] += "\n"
                        current_day_data["taetigkeiten"] += para.text.strip()
            
            # Letzten erfassten Tag hinzufügen
            if current_day_data:
                tage_daten.append(current_day_data)

            # Sortieren und mit leeren Tagen auffüllen
            sorted_tage = {tag_data['tag_name']: tag_data for tag_data in tage_daten}
            final_tage_daten = []
            for day_name in config.DAYS_IN_WEEK:
                final_tage_daten.append(sorted_tage.get(day_name, {
                    "typ": "-", "stunden": "0:00", "taetigkeiten": "-"
                }))

            context["tage_daten"] = final_tage_daten
            
            if not any(day['taetigkeiten'] and day['taetigkeiten'] != '-' for day in context["tage_daten"]):
                 return None, "Keine Tätigkeiten im Bericht gefunden."

            logger.info(f"Datei '{os.path.basename(file_path)}' erfolgreich analysiert.")
            return context, None

        except Exception as e:
            logger.error(f"Fehler beim Analysieren der Datei '{os.path.basename(file_path)}': {e}", exc_info=True)
            return None, f"Allgemeiner Fehler: {e}"
        